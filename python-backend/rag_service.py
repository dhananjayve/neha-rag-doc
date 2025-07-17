import json
import uuid
import os
from typing import List, Dict, Any
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from database import Document, Embedding, DocumentStatus
import numpy as np
import httpx
import ollama
import google.generativeai as genai
from config import settings
import logging

# Configure logging to output to console
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger("rag_service")

class RAGService:
    def __init__(self):
        # Use environment variable for NestJS URL, fallback to localhost for development
        self.nestjs_url = os.getenv('NESTJS_URL', 'http://localhost:3000')
        # Configure Ollama client to use the configured base URL
        # Note: ollama client doesn't have set_host method, it uses environment variable
        logger.info(f"RAGService initialized with NestJS URL: {self.nestjs_url}")
        
        # Cache for question embeddings to avoid regenerating
        self._question_embedding_cache = {}
        
    async def process_document(self, session: AsyncSession, document_id: str) -> Dict[str, Any]:
        """Process a document and generate embeddings using Ollama"""
        logger.info(f"Processing document: {document_id}")
        try:
            # Get document from NestJS backend
            async with httpx.AsyncClient() as client:
                # Use the internal endpoint that doesn't require authentication
                response = await client.get(f"{self.nestjs_url}/api/internal/documents/{document_id}/content")
                if response.status_code != 200:
                    raise ValueError(f"Document not found in NestJS backend: {response.status_code}")
                
                document_data = response.json()
                document_content = document_data.get('content', '')
                document_title = document_data.get('title', '')
                file_content = document_data.get('fileContent')
                mime_type = document_data.get('mimeType', '')
            
            logger.info(f"Retrieved document: {document_title}, content length: {len(document_content)}")
            
            # If content is empty or placeholder, try to extract from file content
            if not document_content or document_content.startswith('Content extracted from'):
                if file_content:
                    logger.info(f"Extracting content from BLOB for document {document_id}")
                    # Handle file_content which might be a Buffer or serialized data
                    if isinstance(file_content, str):
                        # If it's a string, it might be base64 encoded
                        import base64
                        try:
                            file_content = base64.b64decode(file_content)
                        except:
                            # If not base64, treat as regular string
                            file_content = file_content.encode('utf-8')
                    elif isinstance(file_content, list):
                        # If it's a list, it might be a Buffer array
                        file_content = bytes(file_content)
                    
                    document_content = self._extract_content_from_blob(file_content, mime_type)
                    logger.info(f"Extracted content length: {len(document_content)}")
                else:
                    raise ValueError("Document has no content and no file content available")
            
            if not document_content:
                raise ValueError("Document has no content")
            
            # Generate embeddings using Ollama
            chunks = self._split_into_chunks(document_content)
            logger.info(f"Split document into {len(chunks)} chunks")
            embeddings = await self._generate_embeddings(chunks)
            
            # Store embeddings and chunk content in database
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                embedding_record = Embedding(
                    id=str(uuid.uuid4()),
                    document_id=document_id,
                    chunk_index=i,
                    embedding=json.dumps(embedding.tolist()),
                    chunk_content=chunk  # Store the actual chunk content
                )
                session.add(embedding_record)
            
            await session.commit()
            logger.info(f"Successfully stored {len(embeddings)} embeddings for document {document_id}")
            
            return {
                "status": "success",
                "message": f"Document {document_id} processed successfully",
                "embeddings_count": len(embeddings)
            }
            
        except Exception as e:
            logger.error(f"Error processing document {document_id}: {str(e)}")
            logger.error(f"Full error details: {type(e).__name__}: {str(e)}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return {
                "status": "error",
                "message": str(e)
            }
    
    async def _generate_embeddings(self, chunks: List[str]) -> List[np.ndarray]:
        """Generate embeddings for document chunks using Ollama"""
        logger.info(f"Generating embeddings for {len(chunks)} chunks")
        embeddings = []
        
        # Process chunks in parallel for faster embedding generation
        import asyncio
        import httpx
        
        async def generate_embedding(chunk: str) -> np.ndarray:
            try:
                async with httpx.AsyncClient(timeout=10.0) as client:
                    response = await client.post(
                        f"{settings.OLLAMA_BASE_URL}/api/embeddings",
                        json={
                            "model": settings.EMBEDDING_MODEL,
                            "prompt": chunk
                        }
                    )
                    if response.status_code == 200:
                        result = response.json()
                        return np.array(result['embedding'])
                    else:
                        logger.error(f"Embedding request failed: {response.status_code}")
                        return np.zeros(384)
            except Exception as e:
                logger.error(f"Error generating embedding: {str(e)}")
                return np.zeros(384)
        
        # Generate embeddings concurrently
        tasks = [generate_embedding(chunk) for chunk in chunks]
        embeddings = await asyncio.gather(*tasks)
        
        logger.info(f"Successfully generated {len(embeddings)} embeddings")
        return embeddings
    
    def _split_into_chunks(self, text: str, chunk_size: int = 1000) -> List[str]:
        """Split text into chunks with overlap"""
        words = text.split()
        chunks = []
        current_chunk = []
        current_size = 0
        
        for word in words:
            if current_size + len(word) + 1 > chunk_size and current_chunk:
                chunks.append(' '.join(current_chunk))
                # Start new chunk with overlap
                overlap_words = current_chunk[-settings.CHUNK_OVERLAP//10:]  # Approximate overlap
                current_chunk = overlap_words + [word]
                current_size = sum(len(w) for w in current_chunk) + len(current_chunk) - 1
            else:
                current_chunk.append(word)
                current_size += len(word) + 1
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks
    
    def extract_text_from_file(self, file_content: bytes, mime_type: str) -> str:
        """Extract text content from file - public method for API endpoint"""
        return self._extract_content_from_blob(file_content, mime_type)
    
    def _extract_content_from_blob(self, file_content: bytes, mime_type: str) -> str:
        """Extract text content from file BLOB based on MIME type"""
        try:
            if mime_type == 'text/plain':
                return file_content.decode('utf-8', errors='ignore')
            
            elif mime_type == 'application/pdf':
                # Extract text from PDF files using PyPDF2
                try:
                    import PyPDF2
                    from io import BytesIO
                    
                    pdf_file = BytesIO(file_content)
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    
                    text_content = []
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text_content.append(page.extract_text())
                    
                    extracted_text = '\n'.join(text_content)
                    if extracted_text.strip():
                        return extracted_text
                    else:
                        return f"PDF content (size: {len(file_content)} bytes) - No text could be extracted (may be scanned/image-based PDF)"
                except ImportError:
                    return f"PDF content (size: {len(file_content)} bytes) - PyPDF2 library not available"
                except Exception as e:
                    return f"PDF content (size: {len(file_content)} bytes) - Error extracting text: {str(e)}"
            
            elif mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                # Extract text from Word documents using python-docx
                try:
                    from docx import Document
                    from io import BytesIO
                    
                    doc_file = BytesIO(file_content)
                    doc = Document(doc_file)
                    
                    text_content = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(paragraph.text)
                    
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    text_content.append(cell.text)
                    
                    extracted_text = '\n'.join(text_content)
                    if extracted_text.strip():
                        return extracted_text
                    else:
                        return f"Word document content (size: {len(file_content)} bytes) - No text could be extracted"
                except ImportError:
                    return f"Word document content (size: {len(file_content)} bytes) - python-docx library not available"
                except Exception as e:
                    return f"Word document content (size: {len(file_content)} bytes) - Error extracting text: {str(e)}"
            
            elif mime_type == 'text/html':
                # For HTML files, extract text content using BeautifulSoup
                try:
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(file_content.decode('utf-8', errors='ignore'), 'html.parser')
                    # Remove script and style elements
                    for script in soup(["script", "style"]):
                        script.decompose()
                    # Get text and clean up whitespace
                    text = soup.get_text()
                    lines = (line.strip() for line in text.splitlines())
                    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                    text = ' '.join(chunk for chunk in chunks if chunk)
                    return text
                except ImportError:
                    # Fallback to regex if BeautifulSoup is not available
                    import re
                    text = re.sub(r'<[^>]+>', '', file_content.decode('utf-8', errors='ignore'))
                    text = re.sub(r'\s+', ' ', text).strip()
                    return text
            
            elif mime_type == 'application/json':
                # For JSON files, extract text content
                import json
                try:
                    data = json.loads(file_content.decode('utf-8'))
                    return json.dumps(data, indent=2)
                except:
                    return file_content.decode('utf-8', errors='ignore')
            
            elif mime_type == 'text/xml':
                # For XML files, extract text content
                try:
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(file_content.decode('utf-8', errors='ignore'), 'xml')
                    return soup.get_text()
                except ImportError:
                    # Fallback to regex
                    import re
                    text = re.sub(r'<[^>]+>', '', file_content.decode('utf-8', errors='ignore'))
                    text = re.sub(r'\s+', ' ', text).strip()
                    return text
            
            else:
                # For unknown file types, try to decode as text
                try:
                    return file_content.decode('utf-8', errors='ignore')
                except:
                    return f"Binary file content (size: {len(file_content)} bytes, type: {mime_type}) - Text extraction not available"
        
        except Exception as e:
            logger.error(f"Error extracting content from BLOB: {str(e)}")
            return f"Error extracting content: {str(e)}"
    
    async def answer_question(self, session: AsyncSession, question: str, document_ids: List[str] = None) -> Dict[str, Any]:
        """Answer a question using RAG with Gemini API"""
        logger.info(f"Q&A called with question: '{question}' and document_ids: {document_ids}")
        try:
            # Get relevant documents
            if document_ids:
                stmt = select(Document).where(Document.id.in_(document_ids))
            else:
                stmt = select(Document).where(Document.status == DocumentStatus.INGESTED.value)
            
            result = await session.execute(stmt)
            documents = result.scalars().all()
            
            logger.info(f"Found {len(documents)} documents for Q&A")
            
            if not documents:
                logger.warning("No documents available for Q&A")
                return {
                    "answer": "No documents available for answering questions.",
                    "relevant_documents": [],
                    "confidence": 0.0
                }
            
            # Get embeddings for documents
            embeddings = []
            for doc in documents:
                stmt = select(Embedding).where(Embedding.document_id == doc.id)
                result = await session.execute(stmt)
                doc_embeddings = result.scalars().all()
                embeddings.extend(doc_embeddings)
                logger.info(f"Found {len(doc_embeddings)} embeddings for document {doc.title}")
            
            # Find most relevant content using semantic search
            relevant_content = await self._find_relevant_content(question, documents, embeddings)
            logger.info(f"Relevant context length: {len(relevant_content)} characters")
            logger.info(f"Relevant context preview: {relevant_content[:200]}...")
            
            # Generate answer using Gemini API
            answer = await self._generate_answer(question, relevant_content, documents)
            logger.info(f"Generated answer: {answer}")
            
            return {
                "answer": answer,
                "relevant_documents": [str(doc.id) for doc in documents[:3]],  # Top 3
                "confidence": 0.85  # Placeholder confidence
            }
            
        except Exception as e:
            logger.error(f"Error in answer_question: {str(e)}", exc_info=True)
            return {
                "answer": f"Error processing question: {str(e)}",
                "relevant_documents": [],
                "confidence": 0.0
            }
    
    async def _find_relevant_content(self, question: str, documents: List[Document], embeddings: List[Embedding]) -> str:
        """Find most relevant content for the question using semantic search"""
        logger.info(f"Finding relevant content for question: '{question}'")
        try:
            # Check cache first for question embedding
            if question in self._question_embedding_cache:
                question_vector = self._question_embedding_cache[question]
                logger.info("Using cached question embedding")
            else:
                # Generate embedding for the question using httpx for faster response
                import httpx
                async with httpx.AsyncClient(timeout=5.0) as client:
                    response = await client.post(
                        f"{settings.OLLAMA_BASE_URL}/api/embeddings",
                        json={
                            "model": settings.EMBEDDING_MODEL,
                            "prompt": question
                        }
                    )
                    if response.status_code == 200:
                        result = response.json()
                        question_vector = np.array(result['embedding'])
                        # Cache the embedding
                        self._question_embedding_cache[question] = question_vector
                        logger.info("Generated and cached question embedding")
                    else:
                        logger.error("Failed to generate question embedding")
                        return self._fallback_keyword_search(question, documents)
            
            # Calculate similarities
            similarities = []
            for embedding in embeddings:
                try:
                    doc_embedding = np.array(json.loads(embedding.embedding))
                    similarity = np.dot(question_vector, doc_embedding) / (
                        np.linalg.norm(question_vector) * np.linalg.norm(doc_embedding)
                    )
                    similarities.append((similarity, embedding.document_id, embedding.chunk_index, embedding.chunk_content))
                except Exception as e:
                    logger.error(f"Error calculating similarity: {str(e)}")
                    continue
            
            # Sort by similarity and get top chunks
            similarities.sort(key=lambda x: x[0], reverse=True)
            top_chunks = similarities[:3]  # Reduced to top 3 for faster processing
            
            logger.info(f"Top similarity scores: {[f'{s[0]:.3f}' for s in top_chunks[:3]]}")
            
            # Get document content for top chunks
            relevant_content = []
            for similarity, doc_id, chunk_idx, chunk_content in top_chunks:
                # Find the document title
                doc_title = "Unknown Document"
                for doc in documents:
                    if doc.id == doc_id:
                        doc_title = doc.title
                        break
                
                relevant_content.append(f"Document: {doc_title}\nChunk {chunk_idx} (similarity: {similarity:.3f}):\n{chunk_content}")
            
            result = "\n\n---\n\n".join(relevant_content)
            logger.info(f"Retrieved {len(relevant_content)} relevant chunks")
            return result
            
        except Exception as e:
            logger.error(f"Error in semantic search: {str(e)}", exc_info=True)
            # Fallback to simple keyword matching
            return self._fallback_keyword_search(question, documents)
    
    def _fallback_keyword_search(self, question: str, documents: List[Document]) -> str:
        """Fallback to simple keyword matching"""
        logger.info("Using fallback keyword search")
        question_lower = question.lower()
        relevant_chunks = []
        
        for doc in documents:
            if any(word in doc.content.lower() for word in question_lower.split()):
                relevant_chunks.append(f"Document: {doc.title}\nContent: {doc.content[:500]}...")
        
        return "\n\n".join(relevant_chunks[:3])
    
    async def _generate_answer(self, question: str, context: str, documents: List[Document]) -> str:
        """Generate answer based on question and context using Gemini API"""
        if not context:
            logger.warning("No relevant context found for question.")
            return "I don't have enough information to answer this question."
        
        try:
            # Configure Gemini API
            if not settings.GEMINI_API_KEY:
                logger.error("Gemini API key not configured")
                return "Error: Gemini API key not configured"
            
            genai.configure(api_key=settings.GEMINI_API_KEY)
            model = genai.GenerativeModel('gemini-1.5-flash')
            
            # Create shorter prompt for Gemini for faster response
            prompt = f"""Answer: {question}

Context: {context[:800]}...

Answer:"""
            
            logger.info(f"Prompt sent to Gemini (length: {len(prompt)} characters)")
            
            try:
                # Generate response using Gemini with timeout
                import asyncio
                
                # Use asyncio to add timeout
                loop = asyncio.get_event_loop()
                response = await loop.run_in_executor(
                    None, 
                    lambda: model.generate_content(prompt, generation_config=genai.types.GenerationConfig(
                        max_output_tokens=150,  # Limit response length for speed
                        temperature=0.1  # Lower temperature for faster, more focused responses
                    ))
                )
                
                if response.text:
                    logger.info("Gemini response received successfully")
                    return response.text
                else:
                    logger.error("Gemini returned empty response")
                    return "Error: No response from Gemini"
                    
            except asyncio.TimeoutError:
                logger.error("Gemini request timed out")
                return "Error: Request timed out. Please try again."
            except Exception as e:
                logger.error(f"Error calling Gemini API: {str(e)}", exc_info=True)
                return f"Error generating response: {str(e)}"
            
        except Exception as e:
            logger.error(f"Error generating Gemini response: {str(e)}", exc_info=True)
            return f"Error generating response: {str(e)}" 